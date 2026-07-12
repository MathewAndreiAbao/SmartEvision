"""Extract header metadata from ILAW-format DLL .docx files (table-based layout).
Generates clean header-only training data for all 3 classifiers.

Each ILAW DLL has:
  Paragraph[0]: "GRADE X DAILY LESSON LOGS | ILAW Format"
  Table 0, Row 1: Learning Area / Asignatura  |  {SUBJECT}
  Table 0, Row 2: Term / Week                  |  {Term X | Week X}
  Table 0, Row 3: Grade Level / Baitang        |  {Grade X}
"""
import os, re, csv
from docx import Document

BASE_DIR = os.path.dirname(__file__)
DATA_DIR = os.path.join(BASE_DIR, 'data')

GRADE_DIRS = {
    'Grade 1': 'grade1_dlls',
    'Grade 2': 'grade2_dlls',
    'Grade 3': 'grade3_dlls',
    'Grade 4': 'grade4_dlls',
    'Grade 5': 'grade5_dlls',
    'Grade 6': 'grade6_dlls',
}

FILE_PATTERN = re.compile(r'^(.+?)_(\d+)_T1W(\d+)', re.I)

SUBJECT_MAP = {
    'AP': 'AP',
    'ENGLISH': 'English',
    'EPP': 'EPP',
    'FILIPINO': 'Filipino',
    'GMRC': 'GMRC',
    'LANGUAGE': 'Language',
    'MAKABANSA': 'Makabansa',
    'MAPEH': 'MAPEH',
    'MAPEH_MUSICARTS': 'MAPEH',
    'MAPEH_PEHEALTH': 'MAPEH',
    'MATH': 'Mathematics',
    'MATHEMATICS': 'Mathematics',
    'MUSICARTS': 'MAPEH',
    'PEHEALTH': 'MAPEH',
    'READING_LITERACY': 'Reading and Literacy',
    'READING AND LITERACY': 'Reading and Literacy',
    'SCIENCE': 'Science',
}


def extract_header_text(docx_path):
    """Extract the header block from an ILAW-format DLL.
    
    Returns a structured dict and a plain-text header string for ML training.
    """
    doc = Document(docx_path)
    
    # Get the first paragraph (title line)
    title_line = ''
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and ('DAILY LESSON LOGS' in t.upper() or 'ILAW' in t.upper()):
            title_line = t
            break
    
    # Get metadata from the first table (rows are fixed in ILAW format)
    subject_raw = ''
    grade_raw = ''
    week_raw = ''
    teacher_raw = ''

    if doc.tables:
        t0 = doc.tables[0]
        rows = t0.rows
        if len(rows) >= 5:
            # Row 1: Learning Area / Asignatura
            subject_raw = rows[1].cells[1].text.strip()
            # Row 2: Term / Week
            week_raw = rows[2].cells[1].text.strip()
            # Row 3: Grade Level
            grade_raw = rows[3].cells[1].text.strip()
            # Row 4: Teacher
            teacher_raw = rows[4].cells[1].text.strip()
    
    # Build structured data
    metadata = {
        'subject_raw': subject_raw,
        'grade_raw': grade_raw,
        'week_raw': week_raw,
        'teacher_raw': teacher_raw,
        'title': title_line,
    }
    
    # Build a clean header text for ML training
    header_parts = [
        title_line,
        f'Learning Area / Asignatura: {subject_raw}' if subject_raw else '',
        f'Term / Week: {week_raw}' if week_raw else '',
        f'Grade Level and Section / Baitang/Pangkat: {grade_raw}' if grade_raw else '',
    ]
    header_text = '\n'.join(p for p in header_parts if p)
    
    return metadata, header_text


def parse_subject(subject_raw):
    """Parse the subject from raw cell text like 'GMRC 1', 'MATHEMATICS 1'."""
    s = subject_raw.strip().upper()
    # Remove trailing grade number (e.g., "GMRC 1" -> "GMRC")
    s = re.sub(r'\s+\d+$', '', s).strip()
    return SUBJECT_MAP.get(s, s.title())


def parse_grade(grade_raw):
    """Parse grade from raw cell text like 'GRADE 1' or 'Grade 1'."""
    g = grade_raw.strip().upper().replace('GRADE', '').replace('GR', '').strip()
    # Remove non-digit chars
    g = re.sub(r'[^1-6K]', '', g)
    if g == 'K':
        return 'Kindergarten'
    if g.isdigit():
        return f'Grade {int(g)}'
    return grade_raw.strip()


def parse_week(week_raw):
    """Parse week number from 'Term 1 | Week 4'."""
    m = re.search(r'WEEK\s*(\d+)', week_raw.upper())
    if m:
        return int(m.group(1))
    return None


def process_all():
    os.makedirs(DATA_DIR, exist_ok=True)
    
    subject_rows = []
    grade_rows = []
    doctype_rows = []
    errors = []
    
    total_files = 0
    
    for grade_label, dirname in GRADE_DIRS.items():
        dll_dir = os.path.join(BASE_DIR, dirname)
        if not os.path.isdir(dll_dir):
            print(f'  SKIP: {dll_dir} not found')
            continue
        
        docx_files = sorted([f for f in os.listdir(dll_dir) if f.endswith('.docx')])
        print(f'\n{grade_label} ({dirname}): {len(docx_files)} files')
        
        for fname in docx_files:
            path = os.path.join(dll_dir, fname)
            
            # Get expected labels from filename
            m = FILE_PATTERN.match(fname)
            if not m:
                errors.append(f'  NO MATCH: {fname}')
                continue
            
            raw_subject = m.group(1)
            grade_num = m.group(2)
            expected_grade = f'Grade {grade_num}'
            expected_subject = SUBJECT_MAP.get(raw_subject.upper(), raw_subject)
            expected_week = int(m.group(3))
            
            # Extract header
            meta, header_text = extract_header_text(path)
            
            if not header_text:
                errors.append(f'  EMPTY: {fname}')
                continue
            
            total_files += 1
            
            # Build training text: the header is our training data
            # For subject: include the full header
            subject_rows.append({'text': header_text, 'label': expected_subject})
            # Also add just the Learning Area line for focused learning
            la_line = [l for l in header_text.split('\n') if 'Learning Area' in l]
            if la_line:
                subject_rows.append({'text': la_line[0], 'label': expected_subject})
            
            # For grade: include the full header
            grade_rows.append({'text': header_text, 'label': expected_grade})
            # Also add just the Grade Level line
            gl_line = [l for l in header_text.split('\n') if 'Grade Level' in l]
            if gl_line:
                grade_rows.append({'text': gl_line[0], 'label': expected_grade})
            # Also add the title line
            if meta['title']:
                grade_rows.append({'text': meta['title'], 'label': expected_grade})
            
            # For doc type: always DLL
            doctype_rows.append({'text': header_text, 'label': 'DLL'})
            if meta['title']:
                doctype_rows.append({'text': meta['title'], 'label': 'DLL'})
    
    # Write CSVs
    def write_csv(filename, rows, label_col='label'):
        path = os.path.join(DATA_DIR, filename)
        with open(path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=['text', label_col])
            writer.writeheader()
            writer.writerows(rows)
        print(f'\n  Written {len(rows)} rows to {filename}')
    
    write_csv('subject_training.csv', subject_rows)
    write_csv('gradelevel_training.csv', grade_rows)
    write_csv('doctype_training.csv', doctype_rows)
    
    # Summary
    from collections import Counter
    print(f'\n{"="*60}')
    print(f'  HEADER-ONLY TRAINING DATA SUMMARY')
    print(f'  Total source files: {total_files}')
    print(f'{"="*60}')
    
    for name, rows in [('Subject', subject_rows), ('Grade Level', grade_rows), ('Doc Type', doctype_rows)]:
        counts = Counter(r['label'] for r in rows)
        print(f'\n  {name}: {len(rows)} rows')
        for k, v in sorted(counts.items(), key=lambda x: -x[1]):
            print(f'    {k}: {v}')
    
    if errors:
        print(f'\n  ERRORS ({len(errors)}):')
        for e in errors:
            print(f'    {e}')


if __name__ == '__main__':
    process_all()
